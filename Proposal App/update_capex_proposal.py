"""
update_capex_proposal.py
========================================================================
FLOW
  Your inputs Excel  →  this script  →  finished CAPEX Word proposal
  (CAPEX_Proposal_        (reads by        (built from the CAPEX
   Inputs.xlsx)          PH_xxx code)      template, blanks + sample
                                           values filled with your data)

Put these THREE files in the SAME folder and double-click this script:
  - CAPEX_Proposal_Inputs.xlsx              (you fill in the yellow cells)
  - CAPEX_Proposal_-_Coding.docx            (the template — do not rename)
  - update_capex_proposal.py                (this file)

Requirements (first run only):  pip install openpyxl python-docx lxml Pillow

HOW INPUTS ARE READ
  Sheet "Proposal Inputs", columns:
     A = Key (PH_001 … PH_165)   ← stable, do NOT edit
     B = Field (human label)
     C = Your Value              ← what you fill in
  The script keys off column A (PH_xxx), so renaming a label in column B
  never breaks the mapping.

NOTES ON THIS CAPEX TEMPLATE (differs from OPEX)
  • Has Project Cost (3.1) and Financial Analysis (3.2) sections instead
    of Solar Tariff, Cost Savings and Termination Charges.
  • Executive Summary Commercial Snapshot: Investment, Payment Terms,
    AMC Cost, Payback Period, Net Savings over 25 Years, CO₂ Reduction.
  • Has additional Plant Layout types: Ground Mounted, Carport/Floating.
  • BOM includes HT Equipment section and Net Metering Cubicle.
  • T&C has payment milestones, AMC, FPEL GST, exclusions.
"""

import os, sys, glob, re, copy
from datetime import datetime, date

try:
    import openpyxl
    from docx import Document
    from docx.shared import Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    os.system(f'"{sys.executable}" -m pip install openpyxl python-docx lxml Pillow')
    import openpyxl
    from docx import Document
    from docx.shared import Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
EXCEL = os.path.join(BASE, "CAPEX_Proposal_Inputs.xlsx")

# ── Behaviour for cells left EMPTY in the Excel ────────────────────────────────
BLANK_MARKER = "XX"
MARK_OPTIONAL_BLANKS = False


# ══════════════════════════════════════════════════════════════════════════════
# 0.  Locate the template
# ══════════════════════════════════════════════════════════════════════════════
def find_template():
    preferred = [
        "CAPEX_Proposal_-_Coding.docx", "CAPEX_Proposal_Template.docx",
        "CAPEX Proposal - Coding.docx", "CAPEX_Proposal.docx",
    ]
    for n in preferred:
        p = os.path.join(BASE, n)
        if os.path.exists(p):
            return p
    for f in glob.glob(os.path.join(BASE, "*.docx")):
        n = os.path.basename(f).lower()
        if n.startswith("proposal_"):
            continue
        if "capex" in n:
            return f
    return None


# ══════════════════════════════════════════════════════════════════════════════
# 1.  Read the Excel inputs  ->  {PH_xxx: value}
# ══════════════════════════════════════════════════════════════════════════════
def group_indian(int_str):
    """'1026277' -> '10,26,277'   (last 3 digits, then pairs)"""
    if len(int_str) <= 3:
        return int_str
    head, tail = int_str[:-3], int_str[-3:]
    parts = []
    while len(head) > 2:
        parts.insert(0, head[-2:])
        head = head[:-2]
    if head:
        parts.insert(0, head)
    return ",".join(parts + [tail])


def group_western(int_str):
    return f"{int(int_str):,}"


def excel_display(value, numfmt):
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, bool):
        return "Yes" if value else "No"
    if isinstance(value, (datetime, date)):
        return value.strftime("%d-%m-%Y")
    if not isinstance(value, (int, float)):
        return str(value).strip()

    fmt = (numfmt or "General").strip()
    if fmt.lower() in ("general", "@", ""):
        if isinstance(value, int):
            return str(value)
        r = float(f"{value:.11g}")
        if r == int(r):
            return str(int(r))
        return f"{r:.11f}".rstrip("0").rstrip(".")

    sections, buf, in_q = [], "", False
    for ch in fmt:
        if ch == '"':
            in_q = not in_q
            buf += ch
        elif ch == ';' and not in_q:
            sections.append(buf); buf = ""
        else:
            buf += ch
    sections.append(buf)

    used_neg_section = False
    if value == 0 and len(sections) >= 3:
        fmt = sections[2]
    elif value < 0 and len(sections) >= 2:
        fmt = sections[1]
        used_neg_section = True
    else:
        fmt = sections[0]

    fmt = re.sub(r"\[[^\]]*\]", "", fmt)

    literals = {}
    def _stash(m):
        tok = "\x02" + chr(ord("A") + len(literals)) + "\x02"
        literals[tok] = m.group(1) if m.lastindex else m.group(0)
        return tok
    fmt = re.sub(r'"([^"]*)"', _stash, fmt)
    fmt = re.sub(r'\\(.)', lambda m: _stash(m), fmt)
    fmt = re.sub(r"_.", "", fmt)
    fmt = re.sub(r"\*.", "", fmt)

    is_pct = "%" in fmt
    if is_pct:
        value = value * 100
        fmt = fmt.replace("%", "\x01")

    core = re.search(r"[#0?][#0?,.]*", fmt)
    core_txt = core.group(0) if core else ""
    prefix = fmt[:core.start()] if core else ""
    suffix = fmt[core.end():] if core else fmt

    if not re.search(r"[0#]", core_txt):
        text = (prefix + suffix).replace("\x01", "%")
        for tok, lit in literals.items():
            text = text.replace(tok, lit)
        return text.strip()

    if "." in core_txt:
        dec = len(re.sub(r"[^0#?]", "", core_txt.split(".", 1)[1]))
    else:
        dec = 0
    grouped = "," in core_txt.split(".")[0]

    neg = (value < 0) and not used_neg_section
    num = abs(float(value))
    q = f"{num:.{dec}f}"
    int_part, _, dec_part = q.partition(".")
    if grouped:
        int_part = group_indian(int_part)
    out = int_part + (("." + dec_part) if dec else "")
    if neg:
        out = "-" + out

    text = prefix + out + suffix
    text = text.replace("\x01", "%")
    for tok, lit in literals.items():
        text = text.replace(tok, lit)
    return text.strip()


def read_excel(path):
    wb_v = openpyxl.load_workbook(path, data_only=True)
    wb_f = openpyxl.load_workbook(path, data_only=False)
    ws_v = wb_v["Proposal Inputs"] if "Proposal Inputs" in wb_v.sheetnames else wb_v.active
    ws_f = wb_f["Proposal Inputs"] if "Proposal Inputs" in wb_f.sheetnames else wb_f.active

    data, fmts, stale, mand, blank = {}, {}, [], {}, {}
    for r in range(1, ws_v.max_row + 1):
        key = ws_v.cell(r, 1).value
        if not (isinstance(key, str) and re.fullmatch(r"PH_\d{3}", key.strip())):
            continue
        key = key.strip()
        cell = ws_v.cell(r, 3)
        val = cell.value
        fmts[key] = ws_f.cell(r, 3).number_format or "General"
        mand[key] = str(ws_v.cell(r, 5).value or "").strip().lower() in ("yes", "manual")
        raw_now = ws_f.cell(r, 3).value
        blank[key] = ((val is None or str(val).strip() == "")
                      and (raw_now is None or str(raw_now).strip() == ""))
        if val is None:
            raw = ws_f.cell(r, 3).value
            if isinstance(raw, str) and raw.strip().upper() in ("=TODAY()", "=NOW()"):
                val = date.today()
            elif isinstance(raw, str) and raw.startswith("="):
                stale.append((key, raw))
                val = ""
            else:
                val = raw
        data[key] = val
    data["__fmt__"] = fmts
    data["__mand__"] = mand
    data["__blank__"] = blank
    if stale:
        print("      NOTE: these formula cells have no saved result yet -")
        for k, f in stale[:6]:
            print(f"            {k}  {f}")
        print("            open the Excel, press Ctrl+S, and re-run.")
    return data


def g(data, key, default=""):
    if key not in data:
        return default
    v = data.get(key)
    if v is None or str(v).strip() == "":
        if BLANK_MARKER and (data.get("__blank__") or {}).get(key):
            if MARK_OPTIONAL_BLANKS or (data.get("__mand__") or {}).get(key):
                return BLANK_MARKER
        return default
    return excel_display(v, (data.get("__fmt__") or {}).get(key, "General"))


def capacity_display(data):
    """Return capacity as '<value> kWp' with Indian comma formatting."""
    raw = g(data, "PH_001", "XXX")
    if not raw:
        return "XXX"
    text = str(raw).strip()
    text = re.sub(r"\s*kWp\s*$", "", text, flags=re.I).strip()
    if not text:
        return "XXX"
    # Apply Indian comma formatting if it's a plain number
    text_clean = text.replace(",", "")
    try:
        num = float(text_clean)
        if num == int(num):
            text = group_indian(str(int(num)))
        else:
            int_part, dec_part = f"{num:.2f}".split(".")
            text = group_indian(int_part) + "." + dec_part
    except ValueError:
        pass  # keep original text if not numeric
    return f"{text} kWp"


def capacity_filename(data):
    """Return capacity for filenames without spaces or a duplicate kWp."""
    display = capacity_display(data)
    return re.sub(r"[^0-9A-Za-z.]", "", display) or "XXXkWp"


def has(data, key):
    return bool(g(data, key))


# ══════════════════════════════════════════════════════════════════════════════
# 2.  Text helpers (run-aware, formatting preserved)
# ══════════════════════════════════════════════════════════════════════════════
def iter_paragraphs(container):
    for p in container.paragraphs:
        yield p
    for t in container.tables:
        for row in t.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def replace_in_paragraph(p, old, new):
    if not old or old == new or not p.runs:
        return False
    changed = False
    search_from = 0
    for _ in range(200):
        runs = p.runs
        full = "".join(r.text for r in runs)
        start = full.find(old, search_from)
        if start == -1:
            break
        end = start + len(old)
        idx, spans = 0, []
        for r in runs:
            spans.append((idx, idx + len(r.text), r))
            idx += len(r.text)
        touched = [(s, e, r) for (s, e, r) in spans if e > start and s < end]
        if not touched:
            break
        first_s, _, first_r = touched[0]
        last_s, _, last_r = touched[-1]
        prefix = first_r.text[: start - first_s]
        suffix = last_r.text[end - last_s:]
        if first_r is last_r:
            first_r.text = prefix + new + suffix
        else:
            first_r.text = prefix + new
            for (s, e, r) in touched[1:-1]:
                r.text = ""
            last_r.text = suffix
        changed = True
        search_from = start + len(new)
    return changed


def replace_everywhere(doc, mapping):
    for p in iter_paragraphs(doc):
        for old, new in mapping.items():
            if old and old in "".join(r.text for r in p.runs):
                replace_in_paragraph(p, old, new)


def set_cell_text(cell, value, bold=None):
    value = "" if value is None else str(value)
    para = cell.paragraphs[0]
    tmpl = para.runs[0] if para.runs else None
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    run = para.add_run(value)
    if tmpl is not None:
        run.font.name = tmpl.font.name
        run.font.size = tmpl.font.size
        run.font.bold = tmpl.font.bold if bold is None else bold
        if tmpl.font.color and tmpl.font.color.rgb:
            run.font.color.rgb = tmpl.font.color.rgb
    if bold is not None:
        run.font.bold = bold
    return run


def set_two_lines(cell, line1, line2):
    tmpl = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r._element.getparent().remove(r._element)

    def _add(par, text):
        run = par.add_run(text)
        if tmpl is not None:
            run.font.name = tmpl.font.name
            run.font.size = tmpl.font.size
        return run
    _add(p, line1)
    p2 = cell.add_paragraph()
    _add(p2, line2)


# ══════════════════════════════════════════════════════════════════════════════
# 3.  Table finders
# ══════════════════════════════════════════════════════════════════════════════
def _labels(t, col=0):
    return [t.rows[i].cells[col].text.strip() for i in range(len(t.rows))]

def _all_tables(doc):
    out = []
    def walk(tabs):
        for t in tabs:
            out.append(t)
            for row in t.rows:
                for cell in row.cells:
                    if cell.tables:
                        walk(cell.tables)
    walk(doc.tables)
    return out

def find_table(doc, predicate):
    for t in _all_tables(doc):
        try:
            if predicate(t):
                return t
        except Exception:
            pass
    return None

def row_by_label(table, label, col=0):
    label = label.lower()
    for row in table.rows:
        if row.cells[col].text.strip().lower() == label:
            return row
    for row in table.rows:
        if row.cells[col].text.strip().lower().startswith(label):
            return row
    return None


# ══════════════════════════════════════════════════════════════════════════════
# 4.  Fill the document
# ══════════════════════════════════════════════════════════════════════════════
def fill_document(doc, d):
    T = _all_tables(doc)

    # ---- 4.1 Executive Summary — header paragraphs (customer name + location) ----
    for i, p in enumerate(doc.paragraphs):
        if p.style and p.style.name and p.text.strip() == "Technical Snapshot":
            if i >= 2:
                name_p = doc.paragraphs[i - 2]
                loc_p = doc.paragraphs[i - 1]
                if has(d, "PH_019") and name_p.text.strip():
                    for r in name_p.runs:
                        r.text = ""
                    if name_p.runs:
                        name_p.runs[0].text = g(d, "PH_019")
                    else:
                        name_p.add_run(g(d, "PH_019"))
                if has(d, "PH_020") and loc_p.text.strip():
                    for r in loc_p.runs:
                        r.text = ""
                    if loc_p.runs:
                        loc_p.runs[0].text = g(d, "PH_020")
                    else:
                        loc_p.add_run(g(d, "PH_020"))
            break

    # ---- 4.1 Executive Summary — Technical Snapshot ----
    t = find_table(doc, lambda x: len(x.columns) == 2 and
                    any(c.text.strip() == "Proposed Capacity (kWp)" for row in x.rows for c in row.cells))
    if t:
        m = {
            "Proposed Capacity (kWp)":   capacity_display(d),
            "Installation Technology":   g(d, "PH_002"),
            "Estimated Generation":      g(d, "PH_003"),
            "% Consumption Replaced":    g(d, "PH_004"),
            "Net Metering":              g(d, "PH_005"),
            "Project Timeline":          g(d, "PH_008"),
            "Exclusions":                g(d, "PH_009"),
        }
        for lbl, val in m.items():
            row = row_by_label(t, lbl)
            if row and val:
                set_cell_text(row.cells[1], val)
        # Key Components — two lines
        kc = row_by_label(t, "Key Components")
        if kc and (has(d, "PH_006") or has(d, "PH_007")):
            set_two_lines(kc.cells[1],
                          f"Modules – {g(d,'PH_006')}",
                          f"Inverters – {g(d,'PH_007')}")

    # ---- 4.2 Executive Summary — Commercial Snapshot (CAPEX-specific) ----
    t = find_table(doc, lambda x: len(x.columns) == 2 and
                    any("investment" in c.text.lower() and "rs" in c.text.lower()
                        for row in x.rows for c in row.cells) and
                    any("payback" in c.text.lower() for row in x.rows for c in row.cells))
    if t:
        for row in t.rows:
            lbl = row.cells[0].text.strip().lower()
            if lbl.startswith("investment"):
                if has(d, "PH_010"): set_cell_text(row.cells[1], g(d, "PH_010"))
            elif lbl.startswith("payment terms"):
                if has(d, "PH_011"): set_cell_text(row.cells[1], g(d, "PH_011"))
            elif lbl.startswith("amc cost"):
                if has(d, "PH_012"): set_cell_text(row.cells[1], g(d, "PH_012"))
            elif lbl.startswith("payback"):
                if has(d, "PH_013"): set_cell_text(row.cells[1], g(d, "PH_013"))
            elif "net savings" in lbl:
                if has(d, "PH_014"): set_cell_text(row.cells[1], g(d, "PH_014"))
            elif "co" in lbl and "reduction" in lbl:
                if has(d, "PH_015"): set_cell_text(row.cells[1], g(d, "PH_015"))

    # ---- 4.3 Proposal Details ----
    t = find_table(doc, lambda x: len(x.columns) == 2 and
                    x.rows[0].cells[0].text.strip() == "Submitted by")
    if t:
        for lbl, key in [("Submitted by", "PH_016"), ("Submitted on", "PH_017"),
                         ("Validity", "PH_018")]:
            row = row_by_label(t, lbl)
            if row and has(d, key):
                set_cell_text(row.cells[1], g(d, key))

    # ---- 4.4 Customer Details ----
    t = find_table(doc, lambda x: len(x.columns) == 2 and
                    x.rows[0].cells[0].text.strip() == "Name" and
                    any(c.text.strip() == "Contact Person" for row in x.rows for c in row.cells))
    if t:
        for lbl, key in [("Name", "PH_019"), ("Corporate Address", "PH_020"),
                         ("Address", "PH_020"),
                         ("Contact Person", "PH_021"), ("Mobile", "PH_022"),
                         ("Email", "PH_023")]:
            row = row_by_label(t, lbl)
            if row and has(d, key):
                set_cell_text(row.cells[1], g(d, key))

    # ---- 4.5 Site Details ----
    t = find_table(doc, lambda x: len(x.columns) == 2 and
                    any(c.text.strip() == "Latitude" for row in x.rows for c in row.cells) and
                    any(c.text.strip() == "Shading" for row in x.rows for c in row.cells))
    if t:
        for lbl, key in [("Site Address", "PH_024"), ("Address", "PH_024"),
                         ("Latitude", "PH_025"),
                         ("Longitude", "PH_026"), ("Type of Installation", "PH_027"),
                         ("Type of Roofs", "PH_028"), ("Availability of Water", "PH_029"),
                         ("Shading", "PH_030")]:
            row = row_by_label(t, lbl)
            if row and has(d, key):
                set_cell_text(row.cells[1], g(d, key))

    # ---- 4.6 Capacity Assessment (3-col: Param | kWp | Notes) ----
    t = find_table(doc, lambda x: len(x.columns) == 3 and
                    x.rows[0].cells[0].text.strip() == "Parameters")
    if t:
        cap = {
            "based on shadow-free area":        ("PH_032", "PH_033"),
            "based on electricity consumption": ("PH_034", "PH_035"),
            "based on state regulatory limit":  ("PH_036", "PH_037"),
            "proposed capacity":                ("PH_038", "PH_039"),
        }
        for row in t.rows[1:]:
            key = row.cells[0].text.strip().lower()
            if key in cap:
                kp, kn = cap[key]
                if has(d, kp): set_cell_text(row.cells[1], g(d, kp))
                if has(d, kn): set_cell_text(row.cells[2], g(d, kn))

    # ---- 4.7 Plant Layout (3 rows: Install / Mounting / Evacuation) ----
    t = find_table(doc, lambda x: len(x.columns) == 2 and len(x.rows) == 3 and
                    x.rows[0].cells[0].text.strip() == "Type of Installation" and
                    x.rows[2].cells[0].text.strip() == "Power Evacuation")
    if t:
        for lbl, key in [("Type of Installation", "PH_040"),
                         ("Type of Mounting", "PH_041"),
                         ("Power Evacuation", "PH_042")]:
            row = row_by_label(t, lbl)
            if row and has(d, key):
                set_cell_text(row.cells[1], g(d, key))

    # ---- 4.8 Rooftop layout table (5-col, header without "Railings") ----
    def _is_layout5(x, railings):
        if len(x.columns) != 5 or len(x.rows) < 3:
            return False
        h4 = x.rows[0].cells[4].text.strip().lower()
        h0 = x.rows[0].cells[0].text.strip().lower()
        if h0 != "shed / building":
            return False
        return ("railings" in h4) == railings

    roof = find_table(doc, lambda x: _is_layout5(x, railings=False))
    if roof:
        r1 = ["PH_043", "PH_044", "PH_045", "PH_046", "PH_047"]
        r2 = ["PH_048", "PH_049", "PH_050", "PH_051", "PH_052"]
        for ci, key in enumerate(r1):
            if has(d, key): set_cell_text(roof.rows[1].cells[ci], g(d, key))
        for ci, key in enumerate(r2):
            if has(d, key): set_cell_text(roof.rows[2].cells[ci], g(d, key))
        if has(d, "PH_053"):
            set_cell_text(roof.rows[-1].cells[1], g(d, "PH_053"))

    shed = find_table(doc, lambda x: _is_layout5(x, railings=True))
    if shed:
        r1 = ["PH_054", "PH_055", "PH_056", "PH_057", "PH_058"]
        r2 = ["PH_059", "PH_060", "PH_061", "PH_062", "PH_063"]
        for ci, key in enumerate(r1):
            if has(d, key): set_cell_text(shed.rows[1].cells[ci], g(d, key))
        for ci, key in enumerate(r2):
            if has(d, key): set_cell_text(shed.rows[2].cells[ci], g(d, key))
        if has(d, "PH_064"):
            set_cell_text(shed.rows[-1].cells[1], g(d, "PH_064"))

    # ---- 4.8b Ground Mounted layout (7-row, 2-col) ----
    t = find_table(doc, lambda x: len(x.columns) == 2 and len(x.rows) >= 7 and
                    x.rows[0].cells[0].text.strip() == "Type of Installation" and
                    any("fencing" in r.cells[0].text.lower() for r in x.rows) and
                    any("plant lighting" in r.cells[0].text.lower() for r in x.rows))
    if t:
        ground_map = {
            "type of installation": "PH_065",
            "azimuth (deg)":        "PH_066",
            "tilt (deg)":           "PH_067",
            "type of mounting":     "PH_068",
            "power evacuation":     "PH_069",
            "fencing":              "PH_070",
            "plant lighting":       "PH_071",
        }
        for row in t.rows:
            lbl = row.cells[0].text.strip().lower()
            key = ground_map.get(lbl)
            if key and has(d, key):
                set_cell_text(row.cells[1], g(d, key))

    # ---- 4.8c Carport / Floating layout ----
    # Find the SECOND 7-row install table (after Ground Mounted)
    found_ground = False
    for tbl in _all_tables(doc):
        if len(tbl.columns) == 2 and len(tbl.rows) >= 7:
            if tbl.rows[0].cells[0].text.strip() == "Type of Installation" and \
               any("fencing" in r.cells[0].text.lower() for r in tbl.rows):
                if not found_ground:
                    found_ground = True
                    continue  # skip Ground Mounted (already handled)
                # This is Carport/Floating
                carport_map = {
                    "type of installation": "PH_072",
                    "azimuth (deg)":        "PH_073",
                    "tilt (deg)":           "PH_074",
                    "type of mounting":     "PH_075",
                    "power evacuation":     "PH_076",
                    "fencing":              "PH_077",
                    "plant lighting":       "PH_078",
                }
                for row in tbl.rows:
                    lbl = row.cells[0].text.strip().lower()
                    key = carport_map.get(lbl)
                    if key and has(d, key):
                        set_cell_text(row.cells[1], g(d, key))
                break

    # ---- 4.9 Estimated Generation ----
    t = find_table(doc, lambda x: len(x.columns) == 2 and
                    x.rows[0].cells[0].text.strip() == "Type of System")
    if t:
        for lbl, key in [("Type of System", "PH_079"),
                         ("Estimated Generation", "PH_080"),
                         ("Degradation", "PH_081"),
                         ("Guaranteed Generation", "PH_082"),
                         ("Plant Life", "PH_083")]:
            row = row_by_label(t, lbl)
            if row and has(d, key):
                set_cell_text(row.cells[1], g(d, key))

    # ---- 4.10 Bill of Material ----
    bom_map = {
        ("solar pv module", "type"):              "PH_085",
        ("solar pv module", "make / model"):      "PH_086",
        ("solar pv module", "efficiency"):        "PH_087",
        ("inverter", "type"):                     "PH_088",
        ("inverter", "env. class / location"):    "PH_089",
        ("inverter", "make / model"):             "PH_090",
        ("module mounting structure", "type"):     "PH_091",
        ("module mounting structure", "make"):     "PH_092",
        ("dc cable", "make / size"):              "PH_093",
        ("ac cable", "make / size"):              "PH_094",
        ("cable tray", "type / size"):            "PH_095",
        ("lt switchgear", "type / make / rating"): "PH_096",
        ("spare feeder (if any)", "type / make / rating"): "PH_097",
        ("energy meter + accb", "make / class"):  "PH_098",
        ("lightning arrestor", "type / make"):    "PH_099",
        ("earthing kit", "make / specs"):         "PH_100",
        ("lifeline", "make"):                     "PH_101",
        ("walkways", "type / size"):              "PH_102",
        ("mesh over skylights", "make / specs"):  "PH_103",
        ("railings over roof", "make / specs"):   "PH_104",
        ("ladder", "make / specs"):               "PH_105",
        ("remote monitoring system", "make / specs"): "PH_106",
        ("irradiation sensor", "make / specs"):   "PH_107",
        ("ambient temp. sensor", "make / specs"): "PH_108",
        ("module cleaning (manual)", "make"):     "PH_109",
        ("water pump / tank / meter", "make"):    "PH_110",
        # HT Equipment (CAPEX-specific)
        ("ht / vcb panel", "type / make / rating"): "PH_111",
        ("transformer", "type / make / rating"):  "PH_112",
        ("transmission / termination", "type / length"): "PH_113",
        ("auxiliary power", "transformer / lt / ups"): "PH_114",
        ("mcr room / scada", "standard / make"):  "PH_115",
        ("boundary / lighting / cctv", "standard / make"): "PH_116",
        ("fire fighting equipment", "make"):      "PH_117",
        # Net Metering
        ("net metering hardware", "make / class"): "PH_118",
        ("net metering cubicle", "make"):         "PH_119",
    }
    for t in _all_tables(doc):
        if len(t.columns) == 5 and t.rows[0].cells[0].text.strip() == "Sr." \
           and t.rows[0].cells[3].text.strip() == "Specifications":
            for row in t.rows[1:]:
                eq = row.cells[1].text.strip().lower()
                de = row.cells[2].text.strip().lower()
                key = bom_map.get((eq, de))
                if key and has(d, key):
                    set_cell_text(row.cells[3], g(d, key))

    # ---- 4.11 Project Schedule ----
    # PH_120 = delivery months; fallback to PH_009 (Project Timeline from Technical Snapshot)
    timeline_val = g(d, "PH_120") or g(d, "PH_009") or g(d, "PH_100")
    if timeline_val:
        t = find_table(doc, lambda x: len(x.columns) == 2 and
                       x.rows[0].cells[0].text.strip() == "Delivery & Installation")
        if t:
            tail = "from receipt of advance or handover of clear site (whichever is later)."
            val = timeline_val.strip()
            # If value already contains "month" (e.g. "4 + 1 Month"), use as-is
            if re.search(r"(?i)month", val):
                set_cell_text(t.rows[0].cells[1], f"{val} {tail}")
            else:
                set_cell_text(t.rows[0].cells[1], f"{val} months {tail}")

    # ---- 4.12 Project Cost (3.1) — CAPEX-SPECIFIC ----
    t = find_table(doc, lambda x: len(x.columns) == 3 and
                    x.rows[0].cells[0].text.strip() == "System Description" and
                    "investment per" in x.rows[0].cells[1].text.lower())
    if t:
        for row in t.rows[1:]:
            lbl = row.cells[0].text.strip().lower()
            if "turnkey" in lbl or "epc" in lbl:
                if has(d, "PH_121"): set_cell_text(row.cells[1], g(d, "PH_121"))
                if has(d, "PH_122"): set_cell_text(row.cells[2], g(d, "PH_122"))
            elif "net metering" in lbl:
                if has(d, "PH_123"): set_cell_text(row.cells[1], g(d, "PH_123"))
                if has(d, "PH_124"): set_cell_text(row.cells[2], g(d, "PH_124"))
            elif "generation panel" in lbl:
                if has(d, "PH_125"): set_cell_text(row.cells[1], g(d, "PH_125"))
                if has(d, "PH_126"): set_cell_text(row.cells[2], g(d, "PH_126"))
            elif lbl == "total":
                if has(d, "PH_127"): set_cell_text(row.cells[1], g(d, "PH_127"))
                if has(d, "PH_128"): set_cell_text(row.cells[2], g(d, "PH_128"))

    # ---- 4.13 Financial Analysis (3.2) — CAPEX-SPECIFIC ----
    t = find_table(doc, lambda x: len(x.columns) == 3 and
                    x.rows[0].cells[0].text.strip() == "Item" and
                    x.rows[0].cells[2].text.strip() == "Value")
    if t:
        fin_map = {
            "system size":                       "PH_132",
            "system cost (incl. gst)":           "PH_133",
            "gst input credit":                  "PH_134",
            "net cost to client":                "PH_135",
            "amc cost":                          "PH_136",
            "solar units generated (year 1)":    "PH_137",
            "present power tariff":              "PH_138",
            "avg. eb tariff increase":           "PH_139",
            "savings in year 1 (post-tax)":      "PH_140",
            "payback period":                    "PH_141",
            "project life":                      "PH_142",
            "inverter life":                     "PH_143",
            "net savings over project life":     "PH_144",
            "total units over project life":     "PH_145",
            "project irr":                       "PH_146",
            "project irr –":                     "PH_146",
            "project irr -":                     "PH_146",
            "equity irr":                        "PH_147",
            "equity irr –":                      "PH_147",
            "equity irr -":                      "PH_147",
            "levelised cost of generation":      "PH_148",
        }
        for row in t.rows[1:]:
            lbl = row.cells[0].text.strip().lower()
            # Try exact match first, then startswith
            key = fin_map.get(lbl)
            if not key:
                for fl, fk in fin_map.items():
                    if lbl.startswith(fl):
                        key = fk
                        break
            if key and has(d, key):
                set_cell_text(row.cells[2], g(d, key))

    # ---- 4.14 T&C — AMC, GST, Validity, Exclusions ----
    t = find_table(doc, lambda x: len(x.columns) == 2 and
                    any("validity" in c.text.lower() for row in x.rows for c in row.cells) and
                    any("payment milestone" in c.text.lower() for row in x.rows for c in row.cells))
    if t:
        for row in t.rows:
            lbl = row.cells[0].text.strip().lower()
            if lbl.startswith("validity"):
                if has(d, "PH_149"):
                    set_cell_text(row.cells[1],
                        f"This quotation shall remain valid for {g(d,'PH_149')} days from the date of this proposal.")
            elif lbl.startswith("amc"):
                if has(d, "PH_150"):
                    set_cell_text(row.cells[1],
                        f"Rs. {g(d,'PH_150')} lakhs in Year 1 with 5% p.a. escalation.")
            elif lbl.startswith("fpel gst"):
                if has(d, "PH_151"):
                    set_cell_text(row.cells[1], f"{g(d,'PH_151')} (State Specific)")

    # Exclusions table (separate, 2-col with "Exclusions" header)
    t = find_table(doc, lambda x: len(x.columns) == 2 and
                    x.rows[0].cells[0].text.strip() == "Exclusions")
    if t and has(d, "PH_152"):
        # Build numbered list from comma-separated exclusions
        items = [x.strip() for x in g(d, "PH_152").split(",") if x.strip()]
        if items:
            text = "\n".join(f"{i+1}. {item}" for i, item in enumerate(items))
            set_cell_text(t.rows[0].cells[1], text)


# ══════════════════════════════════════════════════════════════════════════════
# 5.  Whole-document text swaps
# ══════════════════════════════════════════════════════════════════════════════
def _preserve_space(t_el):
    if t_el.text and (t_el.text != t_el.text.strip()):
        t_el.set(qn('xml:space'), 'preserve')


def _sub_everywhere(doc, pattern, repl, limit_cell=None):
    rx = re.compile(pattern)
    n = 0
    scope = iter_paragraphs(limit_cell) if limit_cell is not None else iter_paragraphs(doc)
    for p in scope:
        joined = "".join(r.text for r in p.runs)
        if not joined:
            continue
        m = rx.search(joined)
        if m:
            replace_in_paragraph(p, m.group(0), rx.sub(repl, m.group(0)))
            n += 1
    if limit_cell is None:
        for t in doc.element.iter(qn('w:t')):
            if t.text and rx.search(t.text):
                t.text = rx.sub(repl, t.text)
                _preserve_space(t)
                n += 1
    return n


def do_text_swaps(doc, d):
    # Irradiation footnote
    if has(d, "PH_084"):
        _sub_everywhere(doc, r"(Subject to\s+)(\[\s*\u2022\s*\]|[\d,\.]+)(\s*kWh/m)",
                        lambda m: m.group(1) + g(d, "PH_084") + m.group(3))

    # USD-INR
    if has(d, "PH_130"):
        _sub_everywhere(doc, r"(1\s*USD\s*=\s*INR\s*)([\d,\.]+)",
                        lambda m: m.group(1) + g(d, "PH_130"))

    # "In words: Rupees ..."
    if has(d, "PH_129"):
        words = g(d, "PH_129").strip()
        if not re.match(r"(?i)^rupees\b", words) and words != BLANK_MARKER:
            words = "Rupees " + words
        _sub_everywhere(doc, r"(In words:\s*)(Rupees\s+.*|.*)",
                        lambda m: m.group(1) + words)

    # AMC Cost in note (₹ 650)
    if has(d, "PH_131"):
        _sub_everywhere(doc, r"(₹\s*)([\d,]+)(\s*/[-–])",
                        lambda m: m.group(1) + g(d, "PH_131") + m.group(3))


# ══════════════════════════════════════════════════════════════════════════════
# 6.  Cover-page & Environmental-Impact stat boxes
# ══════════════════════════════════════════════════════════════════════════════
WPNS = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
NUMUNIT = re.compile(r"^\s*([\u20b9$]?\s*[\d.,]+\s*%?)\s*(.*?)\s*$")


def _box_t_elements(txbx):
    return [t for t in txbx.iter(qn('w:t'))]

def _box_text(txbx):
    return "".join(t.text or "" for t in _box_t_elements(txbx)).strip()

def _box_set(txbx, new):
    ts = _box_t_elements(txbx)
    if not ts:
        return
    full = "".join(t.text or "" for t in ts).strip()

    # ── Case 1: Static PREFIX (e.g. Rs symbol in ts[0]) ──────────────────────
    # Box like: ['Rs', '{', '5', '}']
    # Keep ts[0], replace ts[1] with value, blank ts[2+]
    if len(ts) >= 2 and (ts[0].text or "").strip() in ("₹", "Rs.", "Rs", "$"):
        ts[1].text = new
        _preserve_space(ts[1])
        for extra in ts[2:]:
            extra.text = ""
        return

    # ── Case 2: Static SUFFIX (e.g. ' months' after placeholder) ─────────────
    # Box like: ['{', '7', '}', ' ', 'months']
    # Find the { placeholder runs (up to and including '}') and replace only those.
    # Leave any runs after the closing '}' untouched.
    # Strategy: find the run index of the first '}' or the last numeric run of {N}
    close_brace_idx = None
    for i, t in enumerate(ts):
        if "}" in (t.text or ""):
            close_brace_idx = i
            break
    if close_brace_idx is not None and close_brace_idx < len(ts) - 1:
        # There are runs AFTER the closing brace — static suffix. Replace only [0..close]
        ts[0].text = new
        _preserve_space(ts[0])
        for t in ts[1:close_brace_idx + 1]:
            t.text = ""
        # ts[close_brace_idx+1 ..] left untouched (the static suffix)
        return

    # ── Case 3: Default — plain replacement ───────────────────────────────────
    m = NUMUNIT.match(new)
    num, unit = (m.group(1).strip(), m.group(2).strip()) if m else (new, "")
    if len(ts) >= 2 and unit:
        ts[0].text = num + " "
        _preserve_space(ts[0])
        ts[1].text = unit
        for extra in ts[2:]:
            extra.text = ""
    else:
        ts[0].text = new
        _preserve_space(ts[0])
        for extra in ts[1:]:
            extra.text = ""

def _box_replace_span(txbx, pattern, new):
    ts = _box_t_elements(txbx)
    if not ts:
        return False
    full = "".join(t.text or "" for t in ts)
    m = re.search(pattern, full)
    if not m:
        return False
    start, end = m.start(), m.end()
    idx, spans = 0, []
    for t in ts:
        ln = len(t.text or "")
        spans.append((idx, idx + ln, t))
        idx += ln
    touched = [(a, b, t) for (a, b, t) in spans if b > start and a < end]
    if not touched:
        return False
    fa, _, ft = touched[0]
    la, _, lt = touched[-1]
    prefix = (ft.text or "")[: start - fa]
    suffix = (lt.text or "")[end - la:]
    if ft is lt:
        ft.text = prefix + new + suffix
        _preserve_space(ft)
    else:
        ft.text = prefix + new
        _preserve_space(ft)
        for (a, b, t) in touched[1:-1]:
            t.text = ""
        lt.text = suffix
        _preserve_space(lt)
    return True

def _box_pos(txbx):
    for anc in txbx.iterancestors():
        if anc.tag.endswith('}anchor'):
            out = {}
            for d_ in ('positionH', 'positionV'):
                el = anc.find(WPNS + d_)
                if el is not None:
                    o = el.find(WPNS + 'posOffset')
                    if o is not None and o.text:
                        try:
                            out[d_] = int(o.text) / 360000.0
                        except ValueError:
                            pass
            return out.get('positionH'), out.get('positionV')
        if anc.tag.endswith('}shape'):
            st = anc.get('style', '') or ''
            mh = re.search(r'left:([-\d.]+)pt', st)
            mv = re.search(r'top:([-\d.]+)pt', st)
            return (float(mh.group(1)) / 28.35 if mh else None,
                    float(mv.group(1)) / 28.35 if mv else None)
    return None, None


# Cover stat box positions — these will need calibration against the actual
# CAPEX template. Using the OPEX positions as a starting point; adjust TOL
# or coordinates after a visual check.
STAT_SLOTS = [
    # ── Cover: Project Snapshot row 0 ──────────────────────────────────────────
    # {2}  x=8.30cm  y=8.61cm  — Proposed Capacity (from BOM, kWp)
    # Value: "4,952.61 kWp" — _KWP suffix appends " kWp" if not already present
    ( 8.30,  8.61, "PH_001_KWP", "Cover {2}: Proposed Capacity kWp"),
    # ── Cover: Project Snapshot row 1 ──────────────────────────────────────────
    # {3}  x=12.39cm y=8.61cm  — Est. Generation kWh/yr
    (12.39,  8.61, "PH_154", "Cover {3}: Est. Generation"),
    # {4}  x=16.30cm y=8.71cm  — % Consumption Replaced (Section 1 PH_004)
    (16.30,  8.71, "PH_004", "Cover {4}: % Consumption Replaced"),
    # ── Cover: Project Snapshot row 2 ──────────────────────────────────────────
    # {5}  x=8.47cm  y=11.74cm — Investment ₹ Cr (Section 1 PH_010)
    # Box has "₹" baked in as ts[0]; send only the number+unit e.g. "1.24 Cr"
    ( 8.47, 11.74, "PH_010", "Cover {5}: Investment Rs Cr (no Rs prefix)"),
    # {6}  x=12.37cm y=11.70cm — CO2 Reduction / yr
    (12.37, 11.70, "PH_157", "Cover {6}: CO2 Reduction"),
    # {7}  x=16.14cm y=11.75cm — Project Timeline
    # Box content: "{7} months" — send full value e.g. "5 + 1 months"
    # _box_set will write into runs [0..2] and leave ts[3]=" " ts[4]="months" intact
    (16.14, 11.75, "PH_158", "Cover {7}: Project Timeline (full string)"),
    # ── Cover: Environmental cards (bottom strip) ───────────────────────────────
    # {8}  x=1.54cm  y=19.06cm — Coal Saved
    ( 1.54, 19.06, "PH_159", "Cover {8}: Coal Saved"),
    # {9}  x=7.83cm  y=19.04cm — Water Savings (M litres number only)
    ( 7.83, 19.04, "PH_160", "Cover {9}: Water Savings M litres"),
    # {10} x=14.24cm y=19.06cm — Trees Planted
    (14.24, 19.06, "PH_161", "Cover {10}: Trees Planted"),
    # ── Environmental Impact page 2x2 grid ─────────────────────────────────────
    # {11} x=4.62cm  y=11.59cm — Upper-Left:  Coal Conserved
    ( 4.62, 11.59, "PH_162", "Env page {11} UL: Coal Conserved"),
    # {12} x=13.20cm y=11.59cm — Upper-Right: CO2 Avoided  [swapped from Trees]
    (13.20, 11.59, "PH_165", "Env page {12} UR: CO2 Avoided"),
    # {13} x=4.58cm  y=18.61cm — Lower-Left:  Water Conserved (raw litres)
    ( 4.58, 18.61, "PH_164", "Env page {13} LL: Water Conserved"),
    # {14} x=13.12cm y=18.53cm — Lower-Right: Trees Planted [swapped from CO2]
    (13.12, 18.53, "PH_163", "Env page {14} LR: Trees Planted"),
]


def fill_cover_stats(doc, d, verbose=True):
    boxes = []
    for txbx in doc.element.iter(qn('w:txbxContent')):
        txt = _box_text(txbx)
        if not txt or len(txt) > 70:
            continue
        left, top = _box_pos(txbx)
        boxes.append((txbx, txt, left, top))

    cap = ""
    if has(d, "PH_001"):
        cap = re.sub(r"(?i)\s*kwp\s*$", "", g(d, "PH_001")).strip().rstrip(",")
    subtitle = ""
    if has(d, "PH_019") or has(d, "PH_020"):
        subtitle = f"{g(d,'PH_019')}, {g(d,'PH_020')}".strip().strip(",").strip()

    for txbx, txt, left, top in boxes:
        # CAPEX cover: "<capacity> kWp Solar Capex Proposal" or similar
        if cap and re.match(r"^[\d.,]+\s*kWp\s+Solar\s+(Capex|CAPEX)\s+Proposal$", txt, re.I):
            _box_replace_span(txbx, r"^[\d.,]+\s*kWp", f"{cap} kWp")
        elif subtitle and top is not None and 8.5 <= top <= 10.0 \
                and left is not None and left < 4 and "kWp" not in txt:
            _box_set(txbx, subtitle)

    TOL = 0.4
    missed = []
    for lx, ty, key, desc in STAT_SLOTS:
        actual_key = key.replace("_KWP", "") if key.endswith("_KWP") else key
        val = g(d, actual_key)
        if not val:
            continue
        if key.endswith("_KWP") and "kWp" not in val and "kwp" not in val.lower():
            val = val + " kWp"
        hit = False
        for txbx, txt, left, top in boxes:
            if left is None or top is None:
                continue
            if abs(left - lx) <= TOL and abs(top - ty) <= TOL:
                _box_set(txbx, val)
                hit = True
        if not hit:
            missed.append(f"{desc} ({key})")
    if verbose and missed:
        print("      NOTE: these cover/environment boxes were not found -")
        for m in missed:
            print(f"            {m}")
    return True


# ══════════════════════════════════════════════════════════════════════════════
# 6b. Site Google-Map image
# ══════════════════════════════════════════════════════════════════════════════
def insert_image(doc, img_path):
    if not img_path or not os.path.exists(img_path):
        return False
    for t in _all_tables(doc):
        if len(t.rows) == 1 and len(t.columns) == 1 and \
           "google map" in t.rows[0].cells[0].text.strip().lower():
            cell = t.rows[0].cells[0]
            for p in list(cell.paragraphs):
                for r in list(p.runs):
                    r._element.getparent().remove(r._element)
            run = cell.paragraphs[0].add_run()
            try:
                run.add_picture(img_path, width=Inches(6.2))
                return True
            except Exception as e:
                print(f"      (image insert failed: {e})")
                return False
    return False


# ══════════════════════════════════════════════════════════════════════════════
# 7.  Page numbers + highlight helpers
# ══════════════════════════════════════════════════════════════════════════════
def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        if para.text.strip():
            continue
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        fld1 = OxmlElement('w:fldSimple'); fld1.set(qn('w:instr'), 'PAGE')
        run._r.addnext(fld1)


def highlight_missing(doc):
    if not BLANK_MARKER:
        return 0
    n = 0
    def _mark(run_el):
        rpr = run_el.find(qn('w:rPr'))
        if rpr is None:
            rpr = OxmlElement('w:rPr')
            run_el.insert(0, rpr)
        for old in rpr.findall(qn('w:highlight')):
            rpr.remove(old)
        hl = OxmlElement('w:highlight')
        hl.set(qn('w:val'), 'red')
        rpr.append(hl)
    ph_rx = re.compile(r"^[Xx]{2,4}$")
    for p_ in iter_paragraphs(doc):
        for r in p_.runs:
            t = (r.text or "").strip()
            if t and (t == BLANK_MARKER or ph_rx.match(t)):
                if t != BLANK_MARKER:
                    r.text = BLANK_MARKER
                _mark(r._element)
                n += 1
    for t in doc.element.iter(qn('w:t')):
        txt = (t.text or "").strip()
        if txt and (txt == BLANK_MARKER or ph_rx.match(txt)):
            parent = t.getparent()
            if parent is not None and parent.tag == qn('w:r'):
                if txt != BLANK_MARKER:
                    t.text = BLANK_MARKER
                _mark(parent)
                n += 1
    return n


def strip_all_highlights(doc):
    n = 0
    for p in iter_paragraphs(doc):
        for r in p.runs:
            rpr = r._element.find(qn('w:rPr'))
            if rpr is not None:
                hl = rpr.find(qn('w:highlight'))
                if hl is not None:
                    rpr.remove(hl)
                    n += 1
    for rel in doc.element.iter(qn('w:highlight')):
        rel.getparent().remove(rel)
    return n


# ══════════════════════════════════════════════════════════════════════════════
# 8.  MAIN
# ══════════════════════════════════════════════════════════════════════════════
def main():
    tmpl = find_template()
    for path, name in [(EXCEL, "Excel inputs (CAPEX_Proposal_Inputs.xlsx)"),
                       (tmpl,  "Word template (CAPEX_Proposal_-_Coding.docx)")]:
        if not path or not os.path.exists(path):
            print(f"\n[X]  {name} not found in:\n    {BASE}")
            input("\nPress Enter to exit ...")
            sys.exit(1)

    print("\n" + "=" * 58)
    print("  FOURTH PARTNER ENERGY — CAPEX Proposal Generator")
    print("=" * 58)

    print("\n[1/6] Reading Excel inputs ...")
    d = read_excel(EXCEL)
    customer = g(d, "PH_019", "Customer")
    cap = capacity_display(d)
    print(f"      Template : {os.path.basename(tmpl)}")
    print(f"      Customer : {customer}")
    print(f"      Capacity : {cap}")

    CRITICAL = [("PH_001", "Proposed Capacity"), ("PH_121", "EPC Cost per Wp"),
                ("PH_128", "Total Investment"), ("PH_132", "System Size"),
                ("PH_138", "Present Power Tariff")]
    bad = []
    for k, lbl in CRITICAL:
        v = g(d, k)
        if v in ("", "0", "-", BLANK_MARKER) or re.fullmatch(r"0+(\.0+)?", v or ""):
            bad.append(f"{lbl} ({k}) = {v or 'empty'!s}")
    if bad:
        print("      !! CHECK THESE - they are zero or empty:")
        for b in bad:
            print(f"         {b}")

    print("[2/6] Loading template ...")
    doc = Document(tmpl)

    print("[3/6] Filling tables & snapshots ...")
    fill_document(doc, d)

    print("[4/6] Swapping sample text (customer / USD / irradiation) ...")
    do_text_swaps(doc, d)
    fill_cover_stats(doc, d)

    print("[5/6] Inserting site image (if provided) ...")
    insert_image(doc, g(d, "PH_031"))

    print("[6/6] Page numbers + highlight cleanup ...")
    add_page_numbers(doc)
    strip_all_highlights(doc)
    n_missing = highlight_missing(doc)
    if n_missing:
        print(f"      {n_missing} blank input(s) marked '{BLANK_MARKER}' in red")

    date_str = datetime.now().strftime("%d_%b_%Y")
    safe_cap = capacity_filename(d)
    out_name = f"Proposal_CAPEX_{customer.replace(' ', '_')}_{safe_cap}_{date_str}.docx"
    out_path = os.path.join(BASE, out_name)
    doc.save(out_path)
    print(f"\n[OK]  Done!\n      {out_path}\n")
    input("Press Enter to close ...")


if __name__ == "__main__":
    main()
