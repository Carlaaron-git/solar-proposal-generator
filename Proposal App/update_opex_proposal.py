"""
update_opex_proposal.py
========================================================================
FLOW
  Your inputs Excel  →  this script  →  finished OPEX Word proposal
  (OPEX_Proposal_        (reads by        (built from the OPEX BOOT/PPA
   Inputs.xlsx)          PH_xxx code)      template, blanks + sample
                                           values filled with your data)

Put these THREE files in the SAME folder and double-click this script:
  - OPEX_Proposal_Inputs.xlsx              (you fill in the yellow cells)
  - OPEX_Proposal_-_Final.docx             (the template — do not rename)
  - update_opex_proposal.py                (this file)

Requirements (first run only):  pip install openpyxl python-docx lxml Pillow

HOW INPUTS ARE READ
  Sheet "Proposal Inputs", columns:
     A = Key (PH_001 … PH_117)   ← stable, do NOT edit
     B = Field (human label)
     C = Your Value              ← what you fill in
  The script keys off column A (PH_xxx), so renaming a label in column B
  never breaks the mapping.

NOTES ON THIS OPEX TEMPLATE (differs from CAPEX)
  • No CAPEX cost / IRR / payback section.  Instead: a Commercial Offer
    block (Offtaker / Power Producer / BOOT terms), a Proposed Solar
    Tariff table, and an Estimated Cost Savings table.
  • The Executive Summary carries a Technical Snapshot AND a Commercial
    Snapshot (grid rate, solar tariff, % savings, CO₂).
  • The Environmental Impact page (5.0) is a fixed image, so its stat
    boxes cannot be written; CO₂/year still flows into the Exec Summary.
  • The Termination-Charges schedule (3.4) is a standard per-kWp table
    and is left as-is.  Scope of Work (2.5) and Standard T&C are static.
"""

import os, sys, glob, re
from datetime import datetime, date

try:
    import openpyxl
    from docx import Document
    from docx.shared import Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    os.system(f'"{sys.executable}" -m pip install openpyxl python-docx lxml Pillow')
    import openpyxl
    from docx import Document
    from docx.shared import Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
EXCEL = os.path.join(BASE, "OPEX_Proposal_Inputs.xlsx")

# ── Behaviour for cells left EMPTY in the Excel ────────────────────────────────
#  BLANK_MARKER = "XX"  -> write XX into the document, highlighted RED, so a
#                          missing input is impossible to miss.
#  BLANK_MARKER = ""    -> write nothing (clears whatever the template had).
#  MARK_OPTIONAL_BLANKS = False -> only MANDATORY fields (column E = "Yes") get
#                          the marker; optional blanks are left untouched.
#                       = True  -> every blank managed field gets the marker.
BLANK_MARKER = "XX"
MARK_OPTIONAL_BLANKS = False


# ══════════════════════════════════════════════════════════════════════════════
# 0.  Locate the template
# ══════════════════════════════════════════════════════════════════════════════
def find_template():
    preferred = [
        "OPEX_Proposal_-_Final.docx", "OPEX_Proposal_-_Final__1_.docx",
        "OPEX Proposal - Final.docx", "OPEX_Proposal_Template.docx",
    ]
    for n in preferred:
        p = os.path.join(BASE, n)
        if os.path.exists(p):
            return p
    for f in glob.glob(os.path.join(BASE, "*.docx")):
        n = os.path.basename(f).lower()
        if n.startswith("proposal_"):          # a previously generated output
            continue
        if "opex" in n:
            return f
    return None


# ══════════════════════════════════════════════════════════════════════════════
# 1.  Read the Excel inputs  ->  {PH_xxx: value}
# ══════════════════════════════════════════════════════════════════════════════
def group_indian(int_str):
    """'1026277' -> '10,26,277'   (last 3 digits, then pairs)"""
    if len(int_str) <= 3:
        return int_str
    head, tail = int_str[:-3], int_str[-3:]
    parts = []
    while len(head) > 2:
        parts.insert(0, head[-2:])
        head = head[:-2]
    if head:
        parts.insert(0, head)
    return ",".join(parts + [tail])


def group_western(int_str):
    return f"{int(int_str):,}"


def excel_display(value, numfmt):
    """Render `value` the way Excel displays it under `numfmt`.

    Handles the formats used in this sheet:
      "0"            2.2222  -> "2"
      "0.00"         750.2   -> "750.20"
      "#,##0"        1026277 -> "10,26,277"   (Indian grouping)
      '"₹ "0.0" L"'  11.2    -> "₹ 11.2 L"
      "0%"           0.47    -> "47%"
      "General"      2.2222  -> "2.2222"
    """
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, bool):
        return "Yes" if value else "No"
    if isinstance(value, (datetime, date)):
        return value.strftime("%d-%m-%Y")
    if not isinstance(value, (int, float)):
        return str(value).strip()

    fmt = (numfmt or "General").strip()
    if fmt.lower() in ("general", "@", ""):
        # Excel's General format shows ~11 significant digits, which hides binary
        # floating-point noise: =C122-C123 stored as 2.5700000000000003 is
        # displayed as 2.57. Mirror that instead of printing the raw float.
        if isinstance(value, int):
            return str(value)
        r = float(f"{value:.11g}")
        if r == int(r):
            return str(int(r))
        return f"{r:.11f}".rstrip("0").rstrip(".")

    # pick the right section: positive ; negative ; zero ; text
    sections, buf, in_q = [], "", False
    for ch in fmt:
        if ch == '"':
            in_q = not in_q
            buf += ch
        elif ch == ';' and not in_q:
            sections.append(buf); buf = ""
        else:
            buf += ch
    sections.append(buf)

    used_neg_section = False
    if value == 0 and len(sections) >= 3:
        fmt = sections[2]                     # dedicated zero section, e.g. _ * "-"??_
    elif value < 0 and len(sections) >= 2:
        fmt = sections[1]
        used_neg_section = True
    else:
        fmt = sections[0]

    # drop colour / condition / locale blocks
    fmt = re.sub(r"\[[^\]]*\]", "", fmt)

    # pull out literal text: quoted runs and backslash-escapes
    literals = {}
    def _stash(m):
        # token must contain no digits / # / , / . so it can't be mistaken
        # for part of the numeric mask
        tok = "\x02" + chr(ord("A") + len(literals)) + "\x02"
        literals[tok] = m.group(1) if m.lastindex else m.group(0)
        return tok
    fmt = re.sub(r'"([^"]*)"', _stash, fmt)
    fmt = re.sub(r'\\(.)', lambda m: _stash(m), fmt)

    # Accounting-format spacing tokens (these are alignment padding in Excel and
    # must not reach the document):
    #   _x  = leave a gap the width of x      -> drop
    #   *x  = repeat x to fill the cell width -> drop
    fmt = re.sub(r"_.", "", fmt)
    fmt = re.sub(r"\*.", "", fmt)

    is_pct = "%" in fmt
    if is_pct:
        value = value * 100
        fmt = fmt.replace("%", "\x01")            # remember position

    core = re.search(r"[#0?][#0?,.]*", fmt)
    core_txt = core.group(0) if core else ""
    prefix = fmt[:core.start()] if core else ""
    suffix = fmt[core.end():] if core else fmt

    # A mask with no real digit placeholder (e.g. the accounting zero section
    # '"-"??') prints its literal only - no number.
    if not re.search(r"[0#]", core_txt):
        text = (prefix + suffix).replace("\x01", "%")
        for tok, lit in literals.items():
            text = text.replace(tok, lit)
        return text.strip()

    # decimals
    if "." in core_txt:
        dec = len(re.sub(r"[^0#?]", "", core_txt.split(".", 1)[1]))
    else:
        dec = 0
    grouped = "," in core_txt.split(".")[0]

    neg = (value < 0) and not used_neg_section
    num = abs(float(value))
    q = f"{num:.{dec}f}"
    int_part, _, dec_part = q.partition(".")
    if grouped:
        # Indian grouping - this sheet carries INR figures throughout.
        int_part = group_indian(int_part)
    out = int_part + (("." + dec_part) if dec else "")
    if neg:
        out = "-" + out

    text = prefix + out + suffix
    text = text.replace("\x01", "%")
    for tok, lit in literals.items():
        text = text.replace(tok, lit)
    return text.strip()


def read_excel(path):
    wb_v = openpyxl.load_workbook(path, data_only=True)   # cached values (TODAY, refs)
    wb_f = openpyxl.load_workbook(path, data_only=False)  # formulas / raw
    ws_v = wb_v["Proposal Inputs"] if "Proposal Inputs" in wb_v.sheetnames else wb_v.active
    ws_f = wb_f["Proposal Inputs"] if "Proposal Inputs" in wb_f.sheetnames else wb_f.active

    data, fmts, stale, mand, blank = {}, {}, [], {}, {}
    for r in range(1, ws_v.max_row + 1):
        key = ws_v.cell(r, 1).value
        if not (isinstance(key, str) and re.fullmatch(r"PH_\d{3}", key.strip())):
            continue
        key = key.strip()
        cell = ws_v.cell(r, 3)
        val = cell.value
        fmts[key] = ws_f.cell(r, 3).number_format or "General"
        mand[key] = str(ws_v.cell(r, 5).value or "").strip().lower() == "yes"
        raw_now = ws_f.cell(r, 3).value
        blank[key] = ((val is None or str(val).strip() == "")
                      and (raw_now is None or str(raw_now).strip() == ""))
        if val is None:                                    # no cached value
            raw = ws_f.cell(r, 3).value
            if isinstance(raw, str) and raw.strip().upper() in ("=TODAY()", "=NOW()"):
                val = date.today()
            elif isinstance(raw, str) and raw.startswith("="):
                stale.append((key, raw))
                val = ""                                   # unresolved formula -> blank
            else:
                val = raw
        data[key] = val
    data["__fmt__"] = fmts
    data["__mand__"] = mand
    data["__blank__"] = blank
    if stale:
        print("      NOTE: these formula cells have no saved result yet -")
        for k, f in stale[:6]:
            print(f"            {k}  {f}")
        print("            open the Excel, press Ctrl+S, and re-run.")
    return data


def g(data, key, default=""):
    """Value for a PH key, rendered exactly as Excel displays it.
    A cell left empty in the Excel yields BLANK_MARKER (highlighted red later)
    when it is mandatory, so nothing silently keeps the template's old value."""
    if key not in data:
        return default
    v = data.get(key)
    if v is None or str(v).strip() == "":
        if BLANK_MARKER and (data.get("__blank__") or {}).get(key):
            if MARK_OPTIONAL_BLANKS or (data.get("__mand__") or {}).get(key):
                return BLANK_MARKER
        return default
    return excel_display(v, (data.get("__fmt__") or {}).get(key, "General"))


def has(data, key):
    return bool(g(data, key))


# ══════════════════════════════════════════════════════════════════════════════
# 2.  Text helpers (run-aware, formatting preserved)
# ══════════════════════════════════════════════════════════════════════════════
def iter_paragraphs(container):
    """Yield every paragraph in a doc/cell, recursing into (nested) tables."""
    for p in container.paragraphs:
        yield p
    for t in container.tables:
        for row in t.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def replace_in_paragraph(p, old, new):
    """Replace `old` with `new` in a paragraph, safe across split runs.
    Cursor-based so it never re-matches text it just inserted (avoids the
    infinite loop when `new` contains `old`)."""
    if not old or old == new or not p.runs:
        return False
    changed = False
    search_from = 0
    for _ in range(200):                              # hard safety cap
        runs = p.runs
        full = "".join(r.text for r in runs)
        start = full.find(old, search_from)
        if start == -1:
            break
        end = start + len(old)
        idx, spans = 0, []
        for r in runs:
            spans.append((idx, idx + len(r.text), r))
            idx += len(r.text)
        touched = [(s, e, r) for (s, e, r) in spans if e > start and s < end]
        if not touched:
            break
        first_s, _, first_r = touched[0]
        last_s, _, last_r = touched[-1]
        prefix = first_r.text[: start - first_s]
        suffix = last_r.text[end - last_s:]
        if first_r is last_r:
            first_r.text = prefix + new + suffix          # match inside one run
        else:
            first_r.text = prefix + new
            for (s, e, r) in touched[1:-1]:
                r.text = ""
            last_r.text = suffix
        changed = True
        search_from = start + len(new)               # move cursor past inserted text
    return changed


def replace_everywhere(doc, mapping):
    for p in iter_paragraphs(doc):
        for old, new in mapping.items():
            if old and old in "".join(r.text for r in p.runs):
                replace_in_paragraph(p, old, new)


def set_cell_text(cell, value, bold=None):
    """Write `value` into a cell, reusing the cell's base font where possible."""
    value = "" if value is None else str(value)
    para = cell.paragraphs[0]
    # remember a template run's formatting
    tmpl = para.runs[0] if para.runs else None
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    run = para.add_run(value)
    if tmpl is not None:
        run.font.name = tmpl.font.name
        run.font.size = tmpl.font.size
        run.font.bold = tmpl.font.bold if bold is None else bold
        if tmpl.font.color and tmpl.font.color.rgb:
            run.font.color.rgb = tmpl.font.color.rgb
    if bold is not None:
        run.font.bold = bold
    return run


def set_two_lines(cell, line1, line2):
    """For the Key-Components cell: two labelled lines."""
    tmpl = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r._element.getparent().remove(r._element)

    def _add(par, text):
        run = par.add_run(text)
        if tmpl is not None:
            run.font.name = tmpl.font.name
            run.font.size = tmpl.font.size
        return run
    _add(p, line1)
    p2 = cell.add_paragraph()
    _add(p2, line2)


# ══════════════════════════════════════════════════════════════════════════════
# 3.  Table finders (signature-based -> robust to index shifts)
# ══════════════════════════════════════════════════════════════════════════════
def _labels(t, col=0):
    return [t.rows[i].cells[col].text.strip() for i in range(len(t.rows))]

def _all_tables(doc):
    """Flatten top-level + nested tables."""
    out = []
    def walk(tabs):
        for t in tabs:
            out.append(t)
            for row in t.rows:
                for cell in row.cells:
                    if cell.tables:
                        walk(cell.tables)
    walk(doc.tables)
    return out

def find_table(doc, predicate):
    for t in _all_tables(doc):
        try:
            if predicate(t):
                return t
        except Exception:
            pass
    return None

def row_by_label(table, label, col=0):
    label = label.lower()
    for row in table.rows:
        if row.cells[col].text.strip().lower() == label:
            return row
    # startswith fallback
    for row in table.rows:
        if row.cells[col].text.strip().lower().startswith(label):
            return row
    return None


# ══════════════════════════════════════════════════════════════════════════════
# 4.  Fill the document
# ══════════════════════════════════════════════════════════════════════════════
def fill_document(doc, d):
    T = _all_tables(doc)

    # ---- 4.1 Executive Summary — header paragraphs (customer name + location) ----
    # These are standalone paragraphs just before "Technical Snapshot", not in a table.
    for i, p in enumerate(doc.paragraphs):
        if p.style and p.style.name and p.text.strip() == "Technical Snapshot":
            # The two paragraphs immediately before are customer name and location
            if i >= 2:
                name_p = doc.paragraphs[i - 2]
                loc_p = doc.paragraphs[i - 1]
                if has(d, "PH_019") and name_p.text.strip():
                    for r in name_p.runs:
                        r.text = ""
                    if name_p.runs:
                        name_p.runs[0].text = g(d, "PH_019")
                    else:
                        name_p.add_run(g(d, "PH_019"))
                if has(d, "PH_020") and loc_p.text.strip():
                    for r in loc_p.runs:
                        r.text = ""
                    if loc_p.runs:
                        loc_p.runs[0].text = g(d, "PH_020")
                    else:
                        loc_p.add_run(g(d, "PH_020"))
            break

    # ---- 4.1 Executive Summary — Technical Snapshot ----
    t = find_table(doc, lambda x: len(x.columns) == 2 and
                    any(c.text.strip() == "Proposed Capacity (kWp)" for c in x.rows[0].cells + x.rows[-1].cells)
                    or (len(x.columns) == 2 and x.rows[0].cells[0].text.strip() == "Proposed Capacity (kWp)"))
    if t:
        m = {
            "Proposed Capacity (kWp)":   g(d, "PH_001"),
            "Installation Technology":   g(d, "PH_002"),
            "Estimated Generation":      g(d, "PH_003"),
            "Guaranteed Generation":     g(d, "PH_004"),
            "% Consumption Replaced":    g(d, "PH_005"),
            "Net Metering":              g(d, "PH_006"),
            "Project Timeline":          g(d, "PH_009"),
            "Exclusions":                g(d, "PH_010"),
        }
        for lbl, val in m.items():
            row = row_by_label(t, lbl)
            if row and val:
                set_cell_text(row.cells[1], val)
        # Key Components — two lines
        kc = row_by_label(t, "Key Components")
        if kc and (has(d, "PH_007") or has(d, "PH_008")):
            set_two_lines(kc.cells[1],
                          f"Modules – {g(d,'PH_007')}",
                          f"Inverters – {g(d,'PH_008')}")

    # ---- 4.2 Executive Summary — Commercial Snapshot ----
    t = find_table(doc, lambda x: len(x.columns) == 2 and
                    any("% grid consumption replaced" in c.text.lower() for row in x.rows for c in row.cells))
    if t:
        for row in t.rows:
            lbl = row.cells[0].text.strip().lower()
            if lbl.startswith("variable grid rate"):
                if has(d, "PH_011"): set_cell_text(row.cells[1], g(d, "PH_011"))
            elif lbl.startswith("proposed solar tariff"):
                if has(d, "PH_012"): set_cell_text(row.cells[1], g(d, "PH_012"))
            elif lbl.startswith("% savings"):
                if has(d, "PH_013"): set_cell_text(row.cells[1], g(d, "PH_013"))
            elif lbl.startswith("% grid consumption"):
                if has(d, "PH_014"): set_cell_text(row.cells[1], g(d, "PH_014"))
            elif lbl.startswith("reduction in co"):
                if has(d, "PH_015"): set_cell_text(row.cells[1], g(d, "PH_015"))

    # ---- 4.3 Proposal Details ----
    t = find_table(doc, lambda x: len(x.columns) == 2 and
                    x.rows[0].cells[0].text.strip() == "Submitted by")
    if t:
        for lbl, key in [("Submitted by", "PH_016"), ("Submitted on", "PH_017"),
                         ("Validity", "PH_018")]:
            row = row_by_label(t, lbl)
            if row and has(d, key):
                set_cell_text(row.cells[1], g(d, key))

    # ---- 4.4 Customer Details ----
    t = find_table(doc, lambda x: len(x.columns) == 2 and
                    x.rows[0].cells[0].text.strip() == "Name" and
                    any(c.text.strip() == "Contact Person" for row in x.rows for c in row.cells))
    if t:
        for lbl, key in [("Name", "PH_019"), ("Address", "PH_020"),
                         ("Contact Person", "PH_021"), ("Mobile", "PH_022"),
                         ("Email", "PH_023")]:
            row = row_by_label(t, lbl)
            if row and has(d, key):
                set_cell_text(row.cells[1], g(d, key))

    # ---- 4.5 Site Details ----
    t = find_table(doc, lambda x: len(x.columns) == 2 and
                    any(c.text.strip() == "Latitude" for row in x.rows for c in row.cells) and
                    any(c.text.strip() == "Shading" for row in x.rows for c in row.cells))
    if t:
        for lbl, key in [("Address", "PH_024"), ("Latitude", "PH_025"),
                         ("Longitude", "PH_026"), ("Type of Installation", "PH_027"),
                         ("Type of Roofs", "PH_028"), ("Availability of Water", "PH_029"),
                         ("Shading", "PH_030")]:
            row = row_by_label(t, lbl)
            if row and has(d, key):
                set_cell_text(row.cells[1], g(d, key))

    # ---- 4.6 Capacity Assessment (3-col: Param | kWp | Notes) ----
    t = find_table(doc, lambda x: len(x.columns) == 3 and
                    x.rows[0].cells[0].text.strip() == "Parameters")
    if t:
        cap = {
            "based on shadow-free area":        ("PH_032", "PH_033"),
            "based on electricity consumption": ("PH_034", "PH_035"),
            "based on state regulatory limit":  ("PH_036", "PH_037"),
            "proposed capacity":                ("PH_038", "PH_039"),
        }
        for row in t.rows[1:]:
            key = row.cells[0].text.strip().lower()
            if key in cap:
                kp, kn = cap[key]
                if has(d, kp): set_cell_text(row.cells[1], g(d, kp))
                if has(d, kn): set_cell_text(row.cells[2], g(d, kn))

    # ---- 4.7 Plant Layout (exactly 3 rows: Install / Mounting / Evacuation) ----
    t = find_table(doc, lambda x: len(x.columns) == 2 and len(x.rows) == 3 and
                    x.rows[0].cells[0].text.strip() == "Type of Installation" and
                    x.rows[2].cells[0].text.strip() == "Power Evacuation")
    if t:
        for lbl, key in [("Type of Installation", "PH_040"),
                         ("Type of Mounting", "PH_041"),
                         ("Power Evacuation", "PH_042")]:
            row = row_by_label(t, lbl)
            if row and has(d, key):
                set_cell_text(row.cells[1], g(d, key))

    # ---- 4.8 Rooftop layout table (5-col, header without "Railings") ----
    def _is_layout5(x, railings):
        if len(x.columns) != 5 or len(x.rows) < 3:
            return False
        h4 = x.rows[0].cells[4].text.strip().lower()
        h0 = x.rows[0].cells[0].text.strip().lower()
        if h0 != "shed / building":
            return False
        return ("railings" in h4) == railings

    roof = find_table(doc, lambda x: _is_layout5(x, railings=False))
    if roof:
        r1 = ["PH_043", "PH_044", "PH_045", "PH_046", "PH_047"]
        r2 = ["PH_048", "PH_049", "PH_050", "PH_051", "PH_052"]
        for ci, key in enumerate(r1):
            if has(d, key): set_cell_text(roof.rows[1].cells[ci], g(d, key))
        for ci, key in enumerate(r2):
            if has(d, key): set_cell_text(roof.rows[2].cells[ci], g(d, key))
        if has(d, "PH_053"):                       # Total row (last row), capacity col
            set_cell_text(roof.rows[-1].cells[1], g(d, "PH_053"))

    shed = find_table(doc, lambda x: _is_layout5(x, railings=True))
    if shed:
        r1 = ["PH_054", "PH_055", "PH_056", "PH_057", "PH_058"]
        r2 = ["PH_059", "PH_060", "PH_061", "PH_062", "PH_063"]
        for ci, key in enumerate(r1):
            if has(d, key): set_cell_text(shed.rows[1].cells[ci], g(d, key))
        for ci, key in enumerate(r2):
            if has(d, key): set_cell_text(shed.rows[2].cells[ci], g(d, key))
        if has(d, "PH_064"):
            set_cell_text(shed.rows[-1].cells[1], g(d, "PH_064"))

    # ---- 4.9 Estimated Generation ----
    t = find_table(doc, lambda x: len(x.columns) == 2 and
                    x.rows[0].cells[0].text.strip() == "Type of System")
    if t:
        for lbl, key in [("Type of System", "PH_065"),
                         ("Estimated Generation", "PH_066"),
                         ("Degradation", "PH_067"),
                         ("Guaranteed Generation", "PH_068"),
                         ("Plant Life", "PH_069")]:
            row = row_by_label(t, lbl)
            if row and has(d, key):
                set_cell_text(row.cells[1], g(d, key))

    # ---- 4.10 Bill of Material (fill Specifications col by Equipment+Description) ----
    bom_map = {
        ("solar pv module", "type"):              "PH_071",
        ("solar pv module", "make / model"):      "PH_072",
        ("solar pv module", "efficiency"):        "PH_073",
        ("inverter", "type"):                     "PH_074",
        ("inverter", "env. class / location"):    "PH_075",
        ("inverter", "make / model"):             "PH_076",
        ("module mounting structure", "type"):    "PH_077",
        ("module mounting structure", "make"):    "PH_078",
        ("dc cable", "make / size"):              "PH_079",
        ("ac cable", "make / size"):              "PH_080",
        ("cable tray", "type / size"):            "PH_081",
        ("lt switchgear", "type / make / rating"):"PH_082",
        ("spare feeder (if any)", "type / make / rating"): "PH_083",
        ("energy meter + accb", "make / class"):  "PH_084",
        ("lightning arrestor", "type / make"):    "PH_085",
        ("earthing kit", "make / specs"):         "PH_086",
        ("lifeline", "make"):                     "PH_087",
        ("walkways", "type / size"):              "PH_088",
        ("mesh over skylights", "make / specs"):  "PH_089",
        ("railings over roof", "make / specs"):   "PH_090",
        ("ladder", "make / specs"):               "PH_091",
        ("remote monitoring system", "make / specs"): "PH_092",
        ("irradiation sensor", "make / specs"):   "PH_093",
        ("ambient temp. sensor", "make / specs"): "PH_094",
        ("module cleaning (manual)", "make"):     "PH_095",
        ("water pump / tank / meter", "make"):    "PH_096",
        ("net metering hardware", "make / class"):"PH_097",
        ("zero feed in device", "make / class"):  "PH_098",
        ("dg synchronization", "make / class"):   "PH_099",
    }
    for t in _all_tables(doc):
        if len(t.columns) == 5 and t.rows[0].cells[0].text.strip() == "Sr." \
           and t.rows[0].cells[3].text.strip() == "Specifications":
            for row in t.rows[1:]:
                eq = row.cells[1].text.strip().lower()
                de = row.cells[2].text.strip().lower()
                key = bom_map.get((eq, de))
                if key and has(d, key):
                    set_cell_text(row.cells[3], g(d, key))

    # ---- 4.11 Project Schedule (only override if provided) ----
    if has(d, "PH_100"):
        t = find_table(doc, lambda x: len(x.columns) == 2 and
                       x.rows[0].cells[0].text.strip() == "Delivery & Installation")
        if t:
            val = g(d, "PH_100")
            # don't produce "4 + 1 Months months" when the input already says
            # "month(s)" - only append the word when it is absent
            tail = "from receipt of advance or handover of clear site"
            if re.search(r"(?i)month", val):
                set_cell_text(t.rows[0].cells[1], f"{val} {tail}")
            else:
                set_cell_text(t.rows[0].cells[1], f"{val} months {tail}")

    # ---- 4.12 Commercial Offer — Offtaker cell (XXX / XX swaps) ----
    t = find_table(doc, lambda x: len(x.columns) == 2 and
                    x.rows[0].cells[0].text.strip() == "Offtaker")
    if t:
        off_cell = t.rows[0].cells[1]
        for p in off_cell.paragraphs:
            if has(d, "PH_101"):
                replace_in_paragraph(p, "XXX", g(d, "PH_101"))
            if has(d, "PH_102"):
                replace_in_paragraph(p, "registered office at XX",
                                     f"registered office at {g(d,'PH_102')}")
        # Title & Tenure contract term
        if has(d, "PH_103"):
            tt = row_by_label(t, "Title & Tenure")
            if tt:
                for p in tt.cells[1].paragraphs:
                    replace_in_paragraph(p, "for a period of 15 years",
                                         f"for a period of {g(d,'PH_103')} years")

    # ---- 4.13 Proposed Solar Tariff (nested table Period|Rate|Type) ----
    t = find_table(doc, lambda x: len(x.columns) == 3 and len(x.rows) >= 2 and
                   x.rows[0].cells[0].text.strip() == "Period" and
                   x.rows[0].cells[1].text.strip().startswith("Tariff"))
    if t:
        data_row = t.rows[1]
        if has(d, "PH_105"): set_cell_text(data_row.cells[0], g(d, "PH_105"))
        if has(d, "PH_106"): set_cell_text(data_row.cells[1], g(d, "PH_106"))
        if has(d, "PH_107"): set_cell_text(data_row.cells[2], g(d, "PH_107"))

    # ---- 4.14 Estimated Cost Savings (Param | Period/Unit | Value) ----
    t = find_table(doc, lambda x: len(x.columns) == 3 and
                   x.rows[0].cells[0].text.strip() == "Parameter" and
                   x.rows[0].cells[2].text.strip() == "Value")
    if t:
        sav = {
            "variable grid rate":               "PH_109",
            "proposed solar tariff":            "PH_110",
            "landed cost of solar power":       "PH_111",
            "per unit savings":                 "PH_112",
            "estimated annual generation (kwh)":"PH_113",
            "1st yr savings (₹)":               "PH_114",
            "total savings over ppa term (₹)":  "PH_115",
            "total savings in words":           "PH_116",
        }
        for row in t.rows[1:]:
            key = sav.get(row.cells[0].text.strip().lower())
            if key and has(d, key):
                set_cell_text(row.cells[2], g(d, key))

    # ---- 4.15 Lock-in period — default 5 years ----
    _lock = g(d, "PH_104") or "5"
    for p in iter_paragraphs(doc):
        replace_in_paragraph(p, "lock-in for 5 years",
                             f"lock-in for {_lock} years")

    # ---- 4.16 Offtaker = Customer Name ----
    if not has(d, "PH_101") and has(d, "PH_019"):
        d["PH_101"] = g(d, "PH_019")


# ══════════════════════════════════════════════════════════════════════════════
# 5.  Whole-document text swaps (customer header, irradiation, USD, In-words)
# ══════════════════════════════════════════════════════════════════════════════
def _sub_everywhere(doc, pattern, repl, limit_cell=None):
    """Regex-replace across body paragraphs AND text boxes."""
    rx = re.compile(pattern)
    n = 0
    scope = iter_paragraphs(limit_cell) if limit_cell is not None else iter_paragraphs(doc)
    for p in scope:
        joined = "".join(r.text for r in p.runs)
        if not joined:
            continue
        m = rx.search(joined)
        if m:
            replace_in_paragraph(p, m.group(0), rx.sub(repl, m.group(0)))
            n += 1
    if limit_cell is None:
        for t in doc.element.iter(qn('w:t')):
            if t.text and rx.search(t.text):
                t.text = rx.sub(repl, t.text)
                _preserve_space(t)
                n += 1
    return n


def do_text_swaps(doc, d):
    """All anchors are PATTERNS, not literal sample values, so they keep working
    after a generated proposal is saved back as the template."""

    # ---- customer name / site location in body tables is handled by
    #      fill_document(); here we only fix free-text mentions ----

    # irradiation footnote:  "Subject to 1,850 kWh/m2/yr" or "[.]"
    if has(d, "PH_070"):
        _sub_everywhere(doc, r"(Subject to\s+)(\[\s*\u2022\s*\]|[\d,\.]+)(\s*kWh/m)",
                        lambda m: m.group(1) + g(d, "PH_070") + m.group(3))

    # USD-INR assumption:  "exchange rate of 1 USD = INR 96.55"
    if has(d, "PH_108"):
        _sub_everywhere(doc, r"(1\s*USD\s*=\s*INR\s*)([\d,\.]+)",
                        lambda m: m.group(1) + g(d, "PH_108"))

    # contract term:  "for a period of 25 years"
    if has(d, "PH_103"):
        _sub_everywhere(doc, r"(for a period of\s+)(\d+)(\s+years)",
                        lambda m: m.group(1) + g(d, "PH_103") + m.group(3))

    # lock-in:  "lock-in for 5 years"
    if has(d, "PH_104"):
        _sub_everywhere(doc, r"(lock-in for\s+)(\d+)(\s+years)",
                        lambda m: m.group(1) + g(d, "PH_104") + m.group(3))

    # "In words: Rupees ..."  -> replace everything after "Rupees"
    if has(d, "PH_116"):
        words = g(d, "PH_116").strip()
        if not re.match(r"(?i)^rupees\b", words) and words != BLANK_MARKER:
            words = "Rupees " + words
        _sub_everywhere(doc, r"(In words:\s*)(.*)$",
                        lambda m: m.group(1) + words)

    # ---- Offtaker block: name + registered office (scoped to that cell only,
    #      so Fourth Partner's own address is never touched) ----
    t = find_table(doc, lambda x: len(x.columns) == 2 and
                   x.rows[0].cells[0].text.strip() == "Offtaker")
    if t is not None:
        cell = t.rows[0].cells[1]
        for p in cell.paragraphs:
            joined = "".join(r.text for r in p.runs)
            # Replace name: everything before "(hereinafter"
            if has(d, "PH_101") and "(hereinafter" in joined:
                old_name = joined.split("(hereinafter")[0].strip()
                if old_name:
                    replace_in_paragraph(p, old_name, g(d, "PH_101") + " ")
            # Replace office: "registered office at <old>."
            if has(d, "PH_102"):
                joined2 = "".join(r.text for r in p.runs)
                m = re.search(r"registered office at\s+(.+?)\.\s*$", joined2)
                if m:
                    replace_in_paragraph(p, m.group(1), g(d, "PH_102"))


# ══════════════════════════════════════════════════════════════════════════════
# 6.  Cover-page & Environmental-Impact stat boxes
#     Matched by POSITION, so re-saving a generated proposal as the template
#     (which changes every value) does not break them.
# ══════════════════════════════════════════════════════════════════════════════
import copy

WPNS = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"

# split "10.26 L" -> ("10.26", "L") ; "47%" -> ("47%", "") ; "5 months" -> ("5","months")
NUMUNIT = re.compile(r"^\s*([\u20b9$]?\s*[\d.,]+\s*%?)\s*(.*?)\s*$")


def _preserve_space(t_el):
    """Keep leading/trailing spaces in a <w:t> (Word strips them otherwise)."""
    if t_el.text and (t_el.text != t_el.text.strip()):
        t_el.set(qn('xml:space'), 'preserve')


def _box_t_elements(txbx):
    return [t for t in txbx.iter(qn('w:t'))]


def _box_text(txbx):
    return "".join(t.text or "" for t in _box_t_elements(txbx)).strip()


def _box_set(txbx, new):
    """Write `new` into a stat box, preserving the template's run split so the
    big-number / small-unit sizing and the space between them stay intact."""
    ts = _box_t_elements(txbx)
    if not ts:
        return
    m = NUMUNIT.match(new)
    num, unit = (m.group(1).strip(), m.group(2).strip()) if m else (new, "")
    if len(ts) >= 2 and unit:
        ts[0].text = num + " "
        _preserve_space(ts[0])
        ts[1].text = unit
        for extra in ts[2:]:
            extra.text = ""
    else:
        ts[0].text = new
        _preserve_space(ts[0])
        for extra in ts[1:]:
            extra.text = ""


def _box_replace_span(txbx, pattern, new):
    """Replace only the part of a text box matching `pattern`, leaving every
    other run (and therefore its font, size and colour) exactly as it was."""
    ts = _box_t_elements(txbx)
    if not ts:
        return False
    full = "".join(t.text or "" for t in ts)
    m = re.search(pattern, full)
    if not m:
        return False
    start, end = m.start(), m.end()
    idx, spans = 0, []
    for t in ts:
        ln = len(t.text or "")
        spans.append((idx, idx + ln, t))
        idx += ln
    touched = [(a, b, t) for (a, b, t) in spans if b > start and a < end]
    if not touched:
        return False
    fa, _, ft = touched[0]
    la, _, lt = touched[-1]
    prefix = (ft.text or "")[: start - fa]
    suffix = (lt.text or "")[end - la:]
    if ft is lt:
        ft.text = prefix + new + suffix
        _preserve_space(ft)
    else:
        ft.text = prefix + new
        _preserve_space(ft)
        for (a, b, t) in touched[1:-1]:
            t.text = ""
        lt.text = suffix
        _preserve_space(lt)
    return True


def _box_pos(txbx):
    """(left_cm, top_cm) of the text box, or (None, None)."""
    for anc in txbx.iterancestors():
        if anc.tag.endswith('}anchor'):
            out = {}
            for d_ in ('positionH', 'positionV'):
                el = anc.find(WPNS + d_)
                if el is not None:
                    o = el.find(WPNS + 'posOffset')
                    if o is not None and o.text:
                        try:
                            out[d_] = int(o.text) / 360000.0
                        except ValueError:
                            pass
            return out.get('positionH'), out.get('positionV')
        if anc.tag.endswith('}shape'):
            st = anc.get('style', '') or ''
            mh = re.search(r'left:([-\d.]+)pt', st)
            mv = re.search(r'top:([-\d.]+)pt', st)
            return (float(mh.group(1)) / 28.35 if mh else None,
                    float(mv.group(1)) / 28.35 if mv else None)
    return None, None


#   (left_cm, top_cm, PH key, description)
STAT_SLOTS = [
    # ── Cover: Project Snapshot row 1 ───────────────────────────────────────
    # x=8.37,  y=6.90  — Proposed Capacity kWp
    ( 8.37,  6.90, "PH_001_KWP", "Cover: Proposed Capacity"),
    # x=12.21, y=6.93  — Est. Generation
    (12.21,  6.93, "PH_118",     "Cover: Est. Generation"),
    # x=16.16, y=6.86  — % Consumption Replaced (the '42%' box)
    (16.16,  6.86, "PH_167",     "Cover: % Consumption Replaced"),
    # ── Cover: Project Snapshot row 2 ───────────────────────────────────────
    # x=8.60,  y=9.98  — Estimated Savings/yr
    ( 8.60,  9.98, "PH_120",     "Cover: Estimated Savings/yr"),
    # x=12.33, y=9.98  — CO2 Reduction/yr
    (12.33,  9.98, "PH_121",     "Cover: CO2 Reduction/yr"),
    # x=16.14, y=10.04 — Project Timeline
    (16.14, 10.04, "PH_122",     "Cover: Project Timeline"),
    # ── Cover: Environmental cards (bottom strip) ───────────────────────────
    # x=1.63,  y=17.45 — Coal Saved
    ( 1.63, 17.45, "PH_123",     "Cover env card: Coal Saved"),
    # x=7.91,  y=17.42 — Water Savings (plain M number)
    ( 7.91, 17.42, "PH_124",     "Cover env card: Water Savings"),
    # x=14.13, y=17.55 — Trees Planted
    (14.13, 17.55, "PH_125",     "Cover env card: Trees Planted"),
    # ── Environmental Impact page 2x2 grid ─────────────────────────────────
    # x=4.45,  y=6.21  — Upper-Left: Coal Conserved
    ( 4.45,  6.21, "PH_126",     "Env page UL: Coal Conserved"),
    # x=13.02, y=6.30  — Upper-Right: Trees Planted
    (13.02,  6.30, "PH_127",     "Env page UR: Trees Planted"),
    # x=4.46,  y=13.17 — Lower-Left: Water Conserved (raw litres)
    ( 4.46, 13.17, "PH_128",     "Env page LL: Water Conserved"),
    # x=12.86, y=13.17 — Lower-Right: CO2 Avoided
    (12.86, 13.17, "PH_129",     "Env page LR: CO2 Avoided"),
]


def fill_cover_stats(doc, d, verbose=True):
    boxes = []
    for txbx in doc.element.iter(qn('w:txbxContent')):
        txt = _box_text(txbx)
        if not txt or len(txt) > 70:
            continue
        left, top = _box_pos(txbx)
        boxes.append((txbx, txt, left, top))

    cap = ""
    if has(d, "PH_001"):
        cap = re.sub(r"(?i)\s*kwp\s*$", "", g(d, "PH_001")).strip().rstrip(",")
    subtitle = ""
    if has(d, "PH_019") or has(d, "PH_020"):
        subtitle = f"{g(d,'PH_019')}, {g(d,'PH_020')}".strip().strip(",").strip()

    for txbx, txt, left, top in boxes:
        if cap and re.match(r"^[\d.,]+\s*kWp\s+Solar\s+Opex\s+Proposal$", txt, re.I):
            # Only the "<capacity> kWp" part is rewritten. "Solar Opex Proposal"
            # sits in separate runs with its own colour (cream vs white) and is
            # left completely untouched.
            _box_replace_span(txbx, r"^[\d.,]+\s*kWp", f"{cap} kWp")
        elif subtitle and top is not None and 7.5 <= top <= 9.2 \
                and left is not None and left < 4 and "kWp" not in txt:
            _box_set(txbx, subtitle)

    TOL = 0.4
    missed = []
    for lx, ty, key, desc in STAT_SLOTS:
        actual_key = key.replace("_KWP", "") if key.endswith("_KWP") else key
        val = g(d, actual_key)
        if not val:
            continue
        if key.endswith("_KWP") and "kWp" not in val and "kwp" not in val.lower():
            val = val + " kWp"
        hit = False
        for txbx, txt, left, top in boxes:
            if left is None or top is None:
                continue
            if abs(left - lx) <= TOL and abs(top - ty) <= TOL:
                _box_set(txbx, val)
                hit = True
        if not hit:
            missed.append(f"{desc} ({key})")
    if verbose and missed:
        print("      NOTE: these cover/environment boxes were not found -")
        for m in missed:
            print(f"            {m}")
    return True


# ══════════════════════════════════════════════════════════════════════════════
# 6b. Termination Charges schedule (3.4) - PH_130..PH_154 = Year 1..25
# ══════════════════════════════════════════════════════════════════════════════
def fill_termination_charges(doc, d):
    """Write termination charges — exactly project_life rows, not always 25."""
    years = []
    for i in range(25):
        v = g(d, f"PH_{130 + i:03d}")
        if v and v not in ("", "0", BLANK_MARKER):
            years.append((i + 1, v))
        elif years:
            break   # stop at first blank after start = project_life boundary

    if not years:
        return False

    t = find_table(doc, lambda x: len(x.columns) == 2 and len(x.rows) >= 2 and
                   x.rows[0].cells[0].text.strip().lower() == "year" and
                   "termination value" in x.rows[0].cells[1].text.strip().lower())
    if t is None:
        print("      (termination charges table not found — skipped)")
        return False

    n = len(years)
    while len(t.rows) - 1 < n:
        t._tbl.append(copy.deepcopy(t.rows[-1]._tr))
    while len(t.rows) - 1 > n:
        t._tbl.remove(t.rows[-1]._tr)

    for idx, (yr, val) in enumerate(years):
        row = t.rows[idx + 1]
        set_cell_text(row.cells[0], f"Year {yr}")
        set_cell_text(row.cells[1], val)

    print(f"      Termination charges: {n} year rows written")
    return True


# ══════════════════════════════════════════════════════════════════════════════
# 6c. Site Google-Map image
# ══════════════════════════════════════════════════════════════════════════════
def insert_image(doc, img_path):
    if not img_path or not os.path.exists(img_path):
        return False
    for t in _all_tables(doc):
        if len(t.rows) == 1 and len(t.columns) == 1 and \
           "google map" in t.rows[0].cells[0].text.strip().lower():
            cell = t.rows[0].cells[0]
            for p in list(cell.paragraphs):
                for r in list(p.runs):
                    r._element.getparent().remove(r._element)
            run = cell.paragraphs[0].add_run()
            try:
                run.add_picture(img_path, width=Inches(6.2))
                return True
            except Exception as e:
                print(f"      (image insert failed: {e})")
                return False
    return False


# ══════════════════════════════════════════════════════════════════════════════
# 7.  Page numbers + yellow-highlight strip
# ══════════════════════════════════════════════════════════════════════════════
def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        if para.text.strip():
            continue
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        fld1 = OxmlElement('w:fldSimple'); fld1.set(qn('w:instr'), 'PAGE')
        run._r.addnext(fld1)


def highlight_missing(doc):
    """Highlight every BLANK_MARKER ("XX") occurrence in red so missing inputs
    are obvious at a glance."""
    if not BLANK_MARKER:
        return 0
    n = 0
    def _mark(run_el):
        rpr = run_el.find(qn('w:rPr'))
        if rpr is None:
            rpr = OxmlElement('w:rPr')
            run_el.insert(0, rpr)
        for old in rpr.findall(qn('w:highlight')):
            rpr.remove(old)
        hl = OxmlElement('w:highlight')
        hl.set(qn('w:val'), 'red')
        rpr.append(hl)
    # Also catch placeholders the template itself still carries (XX / XXX / xxx)
    # because the matching Excel cell was left empty - normalise them to the
    # marker so every unfilled spot looks the same.
    ph_rx = re.compile(r"^[Xx]{2,4}$")
    for p_ in iter_paragraphs(doc):
        for r in p_.runs:
            t = (r.text or "").strip()
            if t and (t == BLANK_MARKER or ph_rx.match(t)):
                if t != BLANK_MARKER:
                    r.text = BLANK_MARKER
                _mark(r._element)
                n += 1
    # text boxes
    for t in doc.element.iter(qn('w:t')):
        txt = (t.text or "").strip()
        if txt and (txt == BLANK_MARKER or ph_rx.match(txt)):
            parent = t.getparent()
            if parent is not None and parent.tag == qn('w:r'):
                if txt != BLANK_MARKER:
                    t.text = BLANK_MARKER
                _mark(parent)
                n += 1
    return n


def strip_all_highlights(doc):
    """Remove EVERY highlight (yellow, red, green, any) from the final output
    so the delivered document is completely clean."""
    n = 0
    for p in iter_paragraphs(doc):
        for r in p.runs:
            rpr = r._element.find(qn('w:rPr'))
            if rpr is not None:
                hl = rpr.find(qn('w:highlight'))
                if hl is not None:
                    rpr.remove(hl)
                    n += 1
    # also text boxes
    for rel in doc.element.iter(qn('w:highlight')):
        rel.getparent().remove(rel)
    return n


# ══════════════════════════════════════════════════════════════════════════════
# 8.  MAIN
# ══════════════════════════════════════════════════════════════════════════════
def main():
    tmpl = find_template()
    for path, name in [(EXCEL, "Excel inputs (OPEX_Proposal_Inputs.xlsx)"),
                       (tmpl,  "Word template (OPEX_Proposal_-_Final.docx)")]:
        if not path or not os.path.exists(path):
            print(f"\n[X]  {name} not found in:\n    {BASE}")
            input("\nPress Enter to exit ...")
            sys.exit(1)

    print("\n" + "=" * 58)
    print("  FOURTH PARTNER ENERGY — OPEX Proposal Generator")
    print("=" * 58)

    print("\n[1/6] Reading Excel inputs ...")
    d = read_excel(EXCEL)
    customer = g(d, "PH_019", "Customer")
    cap = g(d, "PH_001", "XXX")
    print(f"      Template : {os.path.basename(tmpl)}")
    print(f"      Customer : {customer}")
    print(f"      Capacity : {cap} kWp")

    # sanity check on the figures that must never go out as zero
    CRITICAL = [("PH_001", "Proposed Capacity"), ("PH_106", "Tariff Rate"),
                ("PH_109", "Variable Grid Rate"), ("PH_110", "Proposed Solar Tariff"),
                ("PH_112", "Per Unit Savings"), ("PH_113", "Annual Generation"),
                ("PH_115", "Total Savings over PPA Term")]
    bad = []
    for k, lbl in CRITICAL:
        v = g(d, k)
        if v in ("", "0", "-", BLANK_MARKER) or re.fullmatch(r"0+(\.0+)?", v or ""):
            bad.append(f"{lbl} ({k}) = {v or 'empty'!s}")
    if bad:
        print("      !! CHECK THESE - they are zero or empty:")
        for b in bad:
            print(f"         {b}")

    print("[2/6] Loading template ...")
    doc = Document(tmpl)

    print("[3/6] Filling tables & snapshots ...")
    fill_document(doc, d)

    print("[4/6] Swapping sample text (customer / tariff / irradiation) ...")
    do_text_swaps(doc, d)
    fill_cover_stats(doc, d)
    fill_termination_charges(doc, d)

    print("[5/6] Inserting site image (if provided) ...")
    insert_image(doc, g(d, "PH_031"))

    print("[6/6] Page numbers + highlight cleanup ...")
    add_page_numbers(doc)
    strip_all_highlights(doc)
    n_missing = highlight_missing(doc)
    if n_missing:
        print(f"      {n_missing} blank input(s) marked '{BLANK_MARKER}' in red")

    date_str = datetime.now().strftime("%d_%b_%Y")
    safe_cap = re.sub(r"[^0-9A-Za-z]", "", cap) or "XXX"
    out_name = f"Proposal_OPEX_{customer.replace(' ', '_')}_{safe_cap}kWp_{date_str}.docx"
    out_path = os.path.join(BASE, out_name)
    doc.save(out_path)
    print(f"\n[OK]  Done!\n      {out_path}\n")
    input("Press Enter to close ...")


if __name__ == "__main__":
    main()
